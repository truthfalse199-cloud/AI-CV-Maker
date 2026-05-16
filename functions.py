# functions.py
import re
import pdfplumber
from docx import Document
import google.generativeai as genai
from fpdf import FPDF
import streamlit as st
import fitz
import pytesseract
from PIL import Image
import tempfile
import os
import platform

# ==================== KONFIGURASI API ====================
# Ganti dengan API key Anda
GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel("gemini-3.1-flash-lite")

# ==================== EKSTRAKSI FILE ====================
# Set path Tesseract untuk Windows lokal
if platform.system() == "Windows":
    # Sesuaikan path dengan lokasi instalasi Tesseract di komputer Anda
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def extract_from_pdf(file) -> str:
    """
    Ekstrak teks dari PDF menggunakan 3 metode:
    1. pdfplumber (untuk PDF digital)
    2. PyMuPDF (fallback cepat)
    3. OCR via pytesseract (untuk scanned PDF)
    """
    # Simpan file upload ke temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(file.getbuffer())
        tmp_path = tmp_file.name

    try:
        # ===== METODE 1: pdfplumber =====
        with pdfplumber.open(tmp_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if text.strip():
                return text.strip()

        # ===== METODE 2: PyMuPDF (fitz) =====
        doc = fitz.open(tmp_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        if text.strip():
            return text.strip()

        # ===== METODE 3: OCR untuk scanned PDF =====
        st.info("PDF tidak memiliki teks digital, mencoba OCR (membaca gambar)... Ini mungkin memakan waktu.")
        doc = fitz.open(tmp_path)
        text = ""
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            # Render halaman sebagai gambar dengan resolusi tinggi
            zoom = 2.0
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            # OCR dengan bahasa Indonesia + Inggris
            page_text = pytesseract.image_to_string(img, lang='ind+eng')
            text += page_text + "\n"
        doc.close()
        
        if text.strip():
            return text.strip()
        else:
            return "Error: Tidak dapat mengekstrak teks dari PDF (mungkin file rusak atau password protected)."

    except Exception as e:
        return f"Error ekstraksi PDF: {str(e)}"
    finally:
        # Hapus temporary file
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

def extract_from_docx(file) -> str:
    """Ekstrak teks dari file DOCX menggunakan python-docx"""
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

def extract_from_docx(file) -> str:
    """Ekstrak teks dari file DOCX menggunakan python-docx"""
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

# ==================== PROMPT CONSTANTS ====================
# Prompt ini dipisahkan sebagai konstanta agar mudah dikelola dan diperbarui.

SYSTEM_PROMPT_BUAT_CV = """
Kamu adalah AI Career Consultant profesional yang ahli dalam membuat CV ATS-Friendly 
(Applicant Tracking System) untuk pasar kerja Indonesia.

Tugasmu adalah menyusun CV lengkap, profesional, dan terstruktur berdasarkan data yang diberikan.

ATURAN PEMBUATAN CV:
1. Gunakan Bahasa Indonesia yang formal, lugas, dan profesional
2. Susun konten agar lolos sistem ATS: HINDARI tabel, grafik, dan ikon dekoratif
3. Sertakan kata kunci (keywords) yang relevan dengan posisi/bidang yang dituju
4. Urutan section CV: Profil Singkat → Pengalaman Kerja → Pendidikan → Keahlian → Sertifikasi → Organisasi
5. Deskripsi pengalaman kerja: gunakan bullet dengan kata kerja aktif (Mengelola, Mengembangkan, Memimpin, dsb.)
6. Profil singkat: 3-4 kalimat yang mencerminkan nilai tambah kandidat dan tujuan karirnya
7. Jika ada data yang kurang, isi dengan placeholder dalam tanda kurung siku, misal: [Nama Perusahaan]
8. JANGAN menambahkan informasi yang tidak ada dalam data pengguna
9. Format output: plain text rapi dan terstruktur, siap disalin ke Word/Google Docs

OUTPUT FORMAT:
Tampilkan CV dalam format teks terstruktur dengan section header menggunakan garis pemisah (===).
Contoh:
=== PROFIL SINGKAT ===
...

=== PENGALAMAN KERJA ===
...
""".strip()

# ==================== AI POLISHING ====================
def polish_with_ai(raw_text: str, target_lang: str = "en") -> str:
    """
    Mengubah deskripsi kasar (Bahasa Indonesia/Inggris) menjadi bullet point profesional.
    target_lang: "id" untuk Bahasa Indonesia, "en" untuk Inggris.
    """
    if target_lang == "en":
        prompt = f"""
          Kamu adalah AI Career Editor profesional yang ahli dalam memperbaiki dan mengoptimalkan CV 
        agar memenuhi standar ATS (Applicant Tracking System) dan menarik bagi rekruter di pasar kerja Indonesia.

        Tugasmu adalah memoles CV yang diberikan tanpa mengubah data faktual 
        (nama, tempat, tanggal, nama perusahaan, nama institusi) kecuali diminta secara eksplisit.

        TUGAS UTAMA:
        1. Perbaiki bahasa: gunakan kalimat formal, aktif, dan profesional dalam Bahasa Indonesia
        2. Optimalkan struktur: rapikan urutan section, pastikan konsisten dan logis
        3. Perkuat deskripsi pengalaman: ubah kalimat pasif/lemah menjadi bullet dengan kata kerja aktif yang kuat
        4. Tambahkan kata kunci ATS yang relevan berdasarkan bidang pekerjaan yang terdeteksi
         5. Persingkat kalimat yang terlalu panjang atau bertele-tele
          6. Perbaiki typo, tanda baca, dan kapitalisasi yang tidak konsisten
          7. Pastikan format output adalah plain text ATS-Friendly (TANPA tabel, kolom, atau format kompleks)

          LARANGAN KERAS:
         - JANGAN mengubah nama, angka, tanggal, nama perusahaan, atau nama institusi
         - JANGAN menambahkan informasi baru yang tidak ada dalam CV asli
         - JANGAN menggunakan format markdown (**bold**, # heading) dalam output CV

        Deskripsi: {raw_text}
        Output (Tulis teks CV lengkap yang sudah dioptimalkan dalam bahasa inggris):
        """
    else:
        prompt = f"""
        Kamu adalah AI Career Editor profesional yang ahli dalam memperbaiki dan mengoptimalkan CV 
        agar memenuhi standar ATS (Applicant Tracking System) dan menarik bagi rekruter di pasar kerja Indonesia.

        Tugasmu adalah memoles CV yang diberikan tanpa mengubah data faktual 
        (nama, tempat, tanggal, nama perusahaan, nama institusi) kecuali diminta secara eksplisit.

        TUGAS UTAMA:
        1. Perbaiki bahasa: gunakan kalimat formal, aktif, dan profesional dalam Bahasa Indonesia
        2. Optimalkan struktur: rapikan urutan section, pastikan konsisten dan logis
        3. Perkuat deskripsi pengalaman: ubah kalimat pasif/lemah menjadi bullet dengan kata kerja aktif yang kuat
        4. Tambahkan kata kunci ATS yang relevan berdasarkan bidang pekerjaan yang terdeteksi
         5. Persingkat kalimat yang terlalu panjang atau bertele-tele
          6. Perbaiki typo, tanda baca, dan kapitalisasi yang tidak konsisten
          7. Pastikan format output adalah plain text ATS-Friendly (TANPA tabel, kolom, atau format kompleks)

          LARANGAN KERAS:
         - JANGAN mengubah nama, angka, tanggal, nama perusahaan, atau nama institusi
         - JANGAN menambahkan informasi baru yang tidak ada dalam CV asli
         - JANGAN menggunakan format markdown (**bold**, # heading) dalam output CV

        Deskripsi: {raw_text}
         Output (Tulis teks CV lengkap yang sudah dioptimalkan dalam bahasa Indonesia):
        """
    response = model.generate_content(prompt)
    return response.text.strip()

def generate_professional_summary(user_data: dict, target_lang: str = "en") -> str:
    """Buat ringkasan profil berdasarkan data pengguna (pendidikan, pengalaman, skills)"""
    prompt = f"""
    Buatlah paragraf ringkasan profil  untuk CV dalam bahasa {'Inggris' if target_lang == 'en' else 'Indonesia'}.
    Data pengguna:
    - Pendidikan: {user_data.get('education', 'Tidak disebutkan')}
    - Pengalaman kerja: {user_data.get('work_experience', 'Tidak ada')}
    - Keterampilan: {user_data.get('skills', 'Tidak disebutkan')}
    - Status: {user_data.get('status', 'Pencari kerja')}

    Tulis profesional, fokus pada value yang bisa diberikan.
    """
    response = model.generate_content(prompt)
    return response.text.strip()

# ==================== ATS SCORECARD ====================
def ats_scorecard(resume_text: str) -> dict:
    """
    Memeriksa CV teks dan mengembalikan skor (0-100) serta saran.
    """
    score = 100
    suggestions = []

    # Cek panjang CV (ideal 400-800 kata)
    word_count = len(resume_text.split())
    if word_count < 300:
        score -= 20
        suggestions.append("CV terlalu pendek (<300 kata). Tambahkan lebih banyak detail pencapaian.")
    elif word_count > 1000:
        score -= 10
        suggestions.append("CV terlalu panjang (>1000 kata). Persingkat menjadi 1-2 halaman.")

    # Cek keberadaan action verbs (kata kerja aktif sederhana)
    action_verbs = ["mengelola", "memimpin", "mengembangkan", "meningkatkan", "merancang", 
                    "mengimplementasikan", "mengkoordinasikan", "menyusun", "mencapai"]
    found_verbs = [v for v in action_verbs if v.lower() in resume_text.lower()]
    if len(found_verbs) < 3:
        score -= 15
        suggestions.append("Gunakan lebih banyak kata kerja aktif (contoh: memimpin, mengembangkan, meningkatkan).")

    # Cek keberadaan email (kontak penting)
    if not re.search(r'\b[\w\.-]+@[\w\.-]+\.\w+\b', resume_text):
        score -= 10
        suggestions.append("Email tidak ditemukan. Cantumkan alamat email yang valid.")

    # Cek ada nomor telepon (minimal 10 digit)
    if not re.search(r'\b\d{10,}\b', resume_text):
        score -= 10
        suggestions.append("Nomor telepon tidak ditemukan. Cantumkan nomor yang bisa dihubungi.")

    # Peringatan tentang tabel/grafik (tidak bisa dideteksi dari teks, hanya peringatan umum)
    suggestions.append("Pastikan CV Anda tidak mengandung tabel, grafik, atau kolom ganda (tidak bisa dicek otomatis).")

    score = max(0, score)
    return {"score": score, "suggestions": suggestions}

# ==================== GENERATOR PDF ====================
class ResumePDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 14)
        self.cell(0, 10, 'Curriculum Vitae', 0, 1, 'C')
        self.ln(5)

    def add_section(self, title, content):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, title, 0, 1, 'L')
        self.set_font('Arial', '', 11)
        if isinstance(content, list):
            for line in content:
                self.multi_cell(0, 6, f"• {line}")
        else:
            self.multi_cell(0, 6, content)
        self.ln(3)

def generate_pdf(resume_dict: dict) -> bytes:
    """
    Menghasilkan PDF dari dictionary resume.
    resume_dict = {
        'name': '...',
        'contact': '...',
        'summary': '...',
        'experience': ['bullet1', 'bullet2'],
        'education': '...',
        'skills': '...'
    }
    """
    pdf = ResumePDF()
    pdf.add_page()
    # Nama
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, resume_dict.get('name', 'Nama Tidak Diisi'), 0, 1, 'C')
    pdf.set_font('Arial', '', 10)
    pdf.cell(0, 6, resume_dict.get('contact', ''), 0, 1, 'C')
    pdf.ln(5)

    # Ringkasan
    if resume_dict.get('summary'):
        pdf.add_section('PROFESSIONAL SUMMARY', resume_dict['summary'])

    # Pengalaman
    if resume_dict.get('experience'):
        pdf.add_section('WORK EXPERIENCE', resume_dict['experience'])

    # Pendidikan
    if resume_dict.get('education'):
        pdf.add_section('EDUCATION', resume_dict['education'])

    # Keterampilan
    if resume_dict.get('skills'):
        pdf.add_section('SKILLS', resume_dict['skills'])

    # Output sebagai bytes untuk diunduh
    return pdf.output(dest='S').encode('latin-1')
